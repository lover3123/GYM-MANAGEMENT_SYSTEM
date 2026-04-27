from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "gym management system improved report.docx"


ACCENT = "1F4E5F"
ACCENT_2 = "D9EAF2"
LIGHT = "EEF6F8"
TEXT = RGBColor(31, 41, 55)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="B7C9D3", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips=None):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_twips or 9000))


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width.twips))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, bold=False, size=None, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_center(doc, text, size=12, bold=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    style_run(r, bold=bold, size=size, color=color)
    return p


def add_body(doc, text, style=None, space_after=6, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_after = Pt(space_after)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, bold=True, size=11, color=ACCENT)
        rest = p.add_run(text[len(bold_prefix):])
        style_run(rest, size=11)
    else:
        r = p.add_run(text)
        style_run(r, size=11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(item)
        style_run(r, size=10.5)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        style_run(run, bold=True, size=15 if level == 1 else 12, color=ACCENT)
    return p


def add_callout(doc, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    shade_paragraph(p, LIGHT)
    border_paragraph(p, "9CBBC8")
    r = p.add_run(title)
    style_run(r, bold=True, size=10.5, color=ACCENT)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    shade_paragraph(p2, LIGHT)
    border_paragraph(p2, "9CBBC8", top=False)
    r2 = p2.add_run(body)
    style_run(r2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border_paragraph(paragraph, color="C7CDD1", top=True, bottom=True):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge, enabled in (("top", top), ("bottom", bottom), ("left", True), ("right", True)):
        if not enabled:
            continue
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "3")
        node.set(qn("w:color"), color)


def add_labeled_items(doc, title, rows):
    add_heading(doc, title, 2)
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{label}: ")
        style_run(r, bold=True, size=10.5, color=ACCENT)
        r2 = p.add_run(value)
        style_run(r2, size=10.5)


def add_table(doc, headers, rows, widths=None):
    if widths is None:
        widths = [Inches(6.2 / len(headers)) for _ in headers]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, sum(width.twips for width in widths))
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, ACCENT)
        set_cell_margins(cell)
        set_cell_border(cell, "FFFFFF", "4")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        style_run(r, bold=True, size=9.5, color="FFFFFF")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            if len(rows) > 4 and idx == 0:
                set_cell_shading(cell, "F6FAFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 18 and idx != len(headers) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(text))
            style_run(r, size=9.2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, code):
    lines = code.splitlines()
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0 if i < len(lines) - 1 else 6)
        shade_paragraph(p, "F7F7F7")
        border_paragraph(p, "C7CDD1", top=i == 0, bottom=i == len(lines) - 1)
        r = p.add_run(line or " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.4)
        r.font.color.rgb = RGBColor(22, 78, 99)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def setup_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = TEXT
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Aptos Display" if name == "Title" else "Aptos"
    styles["Title"].font.size = Pt(20)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(ACCENT)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.0)

    add_center(doc, "NITTE MEENAKSHI INSTITUTE OF TECHNOLOGY", 15, True, ACCENT)
    add_center(doc, "(Autonomous Institution affiliated to Visvesvaraya Technological University, Belagavi)", 9.5)
    add_center(doc, "Approved by AICTE and Government of Karnataka", 9.5)
    add_center(doc, "Yelahanka, Bengaluru - 560064", 9.5, space_after=18)
    add_center(doc, "MINI PROJECT REPORT", 16, True, ACCENT, 2)
    add_center(doc, "DATABASE MANAGEMENT SYSTEMS LABORATORY (22CSG44)", 12, True, space_after=16)
    add_center(doc, "on", 10)
    add_center(doc, "GYM MANAGEMENT SYSTEM", 20, True, ACCENT, 16)
    add_center(doc, "Submitted in partial fulfillment of the requirements of Semester 4", 10.5)
    add_center(doc, "Bachelor of Engineering in Computer Science and Engineering", 11, True, space_after=16)

    add_center(doc, "Submitted by", 10.5, True, ACCENT, 6)
    add_center(doc, "Ritik Yadav (1NX24CS064)", 10)
    add_center(doc, "Rohan Rajbanshi (1NX24CS065)", 10)
    add_center(doc, "Suyas Pokharel (1NT23CS248)", 10, space_after=12)
    add_center(doc, "Course Instructor", 10.5, True, ACCENT)
    add_center(doc, "Dr. Kavitha K K", 11, True)
    add_center(doc, "Assistant Professor, Department of Computer Science and Engineering", 10)
    add_center(doc, "Academic Year 2025-2026", 10.5, True, space_after=0)

    page_break(doc)
    add_center(doc, "CERTIFICATE", 16, True, ACCENT, 18)
    add_body(
        doc,
        "This is to certify that the mini project report titled \"Gym Management System\" has been carried out by Ritik Yadav (1NX24CS064), Rohan Rajbanshi (1NX24CS065), and Suyas Pokharel (1NT23CS248), bonafide students of the Department of Computer Science and Engineering, Nitte Meenakshi Institute of Technology, Bengaluru.",
    )
    add_body(
        doc,
        "The work has been submitted in partial fulfillment of the requirements of the Database Management Systems Laboratory (22CSG44) for Semester 4 during the academic year 2025-2026.",
    )
    doc.add_paragraph()
    add_body(doc, "Course Instructor: Dr. Kavitha K K, Assistant Professor, Department of CSE")
    add_body(doc, "Signature: ______________________________")
    add_body(doc, "Head of Department: ______________________________")
    add_body(doc, "Signature: ______________________________")

    page_break(doc)
    add_center(doc, "DECLARATION", 16, True, ACCENT, 18)
    add_body(doc, "We hereby declare that:")
    add_bullets(
        doc,
        [
            "This mini project report is our original work prepared for the Database Management Systems Laboratory.",
            "The work has not been submitted for the award of any degree or examination at any other university, college, or institute.",
            "Any data, diagrams, source code references, or external material used in this report has been acknowledged appropriately.",
            "The report has been prepared with academic integrity and reflects our understanding of database design, implementation, and application development.",
        ],
    )
    doc.add_paragraph()
    add_body(doc, "Ritik Yadav (1NX24CS064)      Signature: ____________________")
    add_body(doc, "Rohan Rajbanshi (1NX24CS065)  Signature: ____________________")
    add_body(doc, "Suyas Pokharel (1NT23CS248)   Signature: ____________________")
    add_body(doc, "Date: 27 April 2026")

    page_break(doc)
    add_center(doc, "ACKNOWLEDGEMENT", 16, True, ACCENT, 18)
    add_body(
        doc,
        "We express our sincere gratitude to Nitte Meenakshi Institute of Technology for providing the academic environment and facilities required to complete this mini project successfully.",
    )
    add_body(
        doc,
        "We thank Dr. Vijaya Shetty S, Head of the Department of Computer Science and Engineering, for her support and encouragement. We are especially grateful to our course instructor, Dr. Kavitha K K, Assistant Professor, Department of CSE, for her guidance, suggestions, and continuous support throughout the development of the Gym Management System.",
    )
    add_body(
        doc,
        "We also thank our classmates, teaching staff, and non-teaching staff for their direct and indirect help during the completion of this project.",
    )

    page_break(doc)
    add_center(doc, "TABLE OF CONTENTS", 16, True, ACCENT, 14)
    toc = [
        "1. Abstract",
        "2. Introduction",
        "3. Problem Statement",
        "4. Objectives of the Work",
        "5. Methodology Followed",
        "6. Results and Discussion",
        "7. Summary",
        "8. References",
    ]
    for item in toc:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item.split(". ", 1)[1])
        style_run(r, size=11)

    page_break(doc)
    add_heading(doc, "1. ABSTRACT")
    add_body(
        doc,
        "The Gym Management System is a database-driven web application designed to digitize the core operations of a fitness center. The system replaces manual registers and scattered spreadsheets with a structured platform for managing members, trainers, membership packages, payments, equipment, and administrative access.",
    )
    add_body(
        doc,
        "The project uses MySQL as the relational database, Node.js with Express.js for the REST API, and a responsive HTML, CSS, and JavaScript frontend. Authentication is implemented using JSON Web Tokens, and passwords are protected using bcrypt hashing. The database schema is normalized to reduce redundancy, maintain referential integrity, and support reliable reporting through views, stored procedures, and triggers.",
    )
    add_body(
        doc,
        "The final implementation demonstrates practical DBMS concepts including entity identification, relationship modelling, primary and foreign keys, constraints, joins, aggregate queries, views, stored procedures, and transaction-oriented design. The system is suitable for small and medium-sized gyms and provides a scalable base for future features such as attendance analytics, renewal reminders, and mobile access.",
    )
    add_callout(
        doc,
        "Project Outcome",
        "A working full-stack Gym Management System with an enhanced MySQL schema, secure backend API, dashboard-ready frontend, and operational reporting queries.",
    )

    add_heading(doc, "2. INTRODUCTION")
    add_body(
        doc,
        "Fitness centers handle a large volume of operational information: member details, membership plans, trainer assignments, payments, equipment records, and renewals. When these tasks are handled manually, the process becomes slow, error-prone, and difficult to audit.",
    )
    add_body(
        doc,
        "The Gym Management System presented in this report centralizes these activities into one web-based application. It demonstrates how a relational database can act as the backbone of a real-world management system while the backend API and frontend interface make the data usable for administrators.",
    )
    add_heading(doc, "2.1 Scope of the System", 2)
    add_bullets(
        doc,
        [
            "Admin login and secure session management.",
            "Member registration, profile management, and membership status tracking.",
            "Trainer management with specialization and salary details.",
            "Membership package creation and enrollment workflows.",
            "Payment recording with monthly revenue reporting.",
            "Equipment inventory tracking with condition status.",
            "Dashboard statistics and reporting endpoints for operational decisions.",
        ],
    )

    add_heading(doc, "3. PROBLEM STATEMENT")
    add_body(
        doc,
        "Many small and medium-sized gyms still depend on paper records, spreadsheets, or disconnected applications. This creates delays in searching records, increases the chance of duplicate or inconsistent data, and makes revenue or renewal tracking difficult.",
    )
    add_heading(doc, "3.1 Key Issues Identified", 2)
    add_bullets(
        doc,
        [
            "Manual member records are difficult to update, search, and validate.",
            "Payment history and revenue summaries are hard to prepare without a centralized database.",
            "Trainer assignments and membership package details are not always connected to member profiles.",
            "Equipment maintenance information may be unavailable when management decisions are needed.",
            "Lack of controlled authentication can expose sensitive member and payment data.",
        ],
    )
    add_heading(doc, "3.2 Proposed Solution", 2)
    add_body(
        doc,
        "The proposed system solves these problems by using a normalized MySQL database connected to a Node.js REST API. The frontend communicates with the API to perform member, trainer, package, membership, payment, and equipment operations. Referential integrity is maintained through foreign keys, and reporting is simplified through database views and aggregate queries.",
    )

    add_heading(doc, "4. OBJECTIVES OF THE WORK")
    add_heading(doc, "4.1 Primary Objectives", 2)
    add_bullets(
        doc,
        [
            "Design a normalized relational schema for a gym management environment.",
            "Implement a RESTful backend API using Node.js and Express.js.",
            "Build a responsive frontend for administrators to manage gym operations.",
            "Apply DBMS concepts such as keys, constraints, joins, views, stored procedures, and triggers.",
        ],
    )
    add_heading(doc, "4.2 Secondary Objectives", 2)
    add_bullets(
        doc,
        [
            "Provide secure admin authentication using JWT and bcrypt.",
            "Generate operational reports such as active memberships and monthly revenue.",
            "Improve maintainability by separating database, backend, and frontend layers.",
            "Create a system that can be extended with attendance tracking, reminders, and analytics.",
        ],
    )

    page_break(doc)
    add_heading(doc, "5. METHODOLOGY FOLLOWED")
    add_body(
        doc,
        "The project followed an iterative software development approach. Requirements were identified first, followed by database design, backend API development, frontend integration, and testing.",
    )
    add_heading(doc, "5.1 Requirements Analysis", 2)
    add_body(
        doc,
        "The main stakeholders are gym administrators, members, and trainers. Administrators require reliable control over member records, packages, payments, and equipment. Members and trainers are represented in the database so their relationships can be tracked accurately.",
    )
    add_heading(doc, "5.2 Database Design", 2)
    add_body(
        doc,
        "The schema was designed using entity-relationship modelling and normalization. The final schema contains seven main operational tables and one admin table. Foreign keys connect memberships to members, packages, and trainers, while payments are linked to both memberships and members.",
    )
    add_labeled_items(
        doc,
        "5.3 Technology Stack",
        [
            ("Database", "MySQL for relational storage, constraints, views, stored procedure, and trigger."),
            ("Backend", "Node.js with Express.js for REST APIs, routing, authentication, and validation."),
            ("Authentication", "JWT for sessions and bcrypt for password hashing."),
            ("Frontend", "HTML5, CSS3, and JavaScript for the browser-based dashboard."),
            ("Visualization", "Chart.js for dashboard charts such as monthly revenue trends."),
            ("Version Control", "Git and GitHub for source code management and collaboration."),
        ],
    )
    add_heading(doc, "5.4 Backend and API Development", 2)
    add_body(
        doc,
        "The backend is organized into controllers, routes, middleware, and database configuration modules. Express.js handles routing, mysql2 manages database access, and a JWT middleware protects authenticated routes.",
    )
    add_bullets(
        doc,
        [
            "POST /api/auth/login - Admin login and JWT generation.",
            "GET /api/members/stats - Dashboard statistics.",
            "GET/POST /api/members - List or add members.",
            "PUT/DELETE /api/members/:id - Update or delete a member.",
            "GET/POST /api/trainers - List or add trainers.",
            "GET/POST /api/packages - List or add membership packages.",
            "GET/POST /api/memberships - List memberships or enroll a member.",
            "GET /api/memberships/expiring - Find memberships expiring within seven days.",
            "GET /api/payments/monthly - Return monthly revenue data for charts.",
        ],
    )
    add_heading(doc, "5.5 Frontend Development", 2)
    add_body(
        doc,
        "The frontend is implemented as a single-page application with a dark dashboard interface. It communicates with backend APIs to display statistics, manage entities, and visualize revenue data.",
    )
    add_heading(doc, "5.6 Testing", 2)
    add_bullets(
        doc,
        [
            "API routes were tested for successful and unsuccessful request flows.",
            "Database constraints were checked through sample inserts and invalid foreign key cases.",
            "Authentication flow was verified using admin credentials and protected routes.",
            "Frontend forms and dashboard data loading were tested against backend responses.",
        ],
    )

    page_break(doc)
    add_heading(doc, "6. RESULTS AND DISCUSSION")
    add_labeled_items(
        doc,
        "6.1 Database Entities",
        [
            ("Admin", "admin_id, username, password, full_name, and email. Stores administrator credentials and account details."),
            ("Member", "member_id, fname, lname, email, contact, gender, join_date, and status. Stores member profiles."),
            ("Trainer", "trainer_id, name, phone, speciality, salary, and status. Maintains trainer information."),
            ("Package", "package_id, package_name, duration_months, and amount. Stores membership packages."),
            ("Membership", "membership_id, member_id, package_id, trainer_id, start_date, end_date, and status. Connects members to packages and trainers."),
            ("Payment", "payment_id, membership_id, member_id, amount, payment_type, and status. Records payment history."),
            ("Equipment", "equipment_id, name, category, last_service, condition_status, and quantity. Tracks gym inventory."),
        ],
    )
    add_heading(doc, "6.2 Entity Relationships", 2)
    add_bullets(
        doc,
        [
            "Member to Membership (1:N): a member can have multiple memberships over time.",
            "Package to Membership (1:N): a package can be selected by many memberships.",
            "Trainer to Membership (1:N): a trainer may be assigned to multiple memberships.",
            "Membership to Payment (1:N): a membership can have one or more related payment records.",
            "Member to Payment (1:N): a member can make multiple payments.",
        ],
    )
    add_heading(doc, "6.3 Schema Highlights", 2)
    add_code_block(
        doc,
        """CREATE TABLE Membership (
  membership_id INT AUTO_INCREMENT PRIMARY KEY,
  member_id INT NOT NULL,
  package_id INT NOT NULL,
  trainer_id INT,
  start_date DATE NOT NULL,
  end_date DATE NOT NULL,
  status ENUM('Active','Expired','Cancelled') DEFAULT 'Active',
  FOREIGN KEY (member_id) REFERENCES Member(member_id) ON DELETE CASCADE,
  FOREIGN KEY (package_id) REFERENCES Package(package_id) ON DELETE RESTRICT,
  FOREIGN KEY (trainer_id) REFERENCES Trainer(trainer_id) ON DELETE SET NULL
);""",
    )
    add_heading(doc, "6.4 Reporting Views and Procedures", 2)
    add_bullets(
        doc,
        [
            "vw_active_memberships returns active members with package fee, dates, and trainer details.",
            "vw_monthly_revenue groups successful payments by month and calculates total revenue.",
            "sp_enroll_member inserts a new member, creates a membership, and records payment in one controlled workflow.",
            "trg_check_membership_expiry updates active memberships to expired when their end date has passed during updates.",
        ],
    )
    add_heading(doc, "6.5 Key SQL Queries", 2)
    add_code_block(
        doc,
        """SELECT member_name, package_name, package_fee, end_date, trainer_name
FROM vw_active_memberships
ORDER BY end_date;

SELECT month, total_payments, total_revenue
FROM vw_monthly_revenue;

SELECT status, COUNT(*) AS member_count
FROM Member
GROUP BY status;""",
    )
    add_heading(doc, "6.6 Discussion", 2)
    add_body(
        doc,
        "The implementation shows that a normalized relational database significantly improves the reliability of gym operations. The use of foreign keys prevents orphaned membership and payment records, while views reduce the complexity of common reporting queries.",
    )
    add_body(
        doc,
        "The stored procedure for enrollment demonstrates how a multi-step workflow can be handled consistently at the database level. The trigger for membership expiry supports data consistency by reducing the need for manual status correction.",
    )

    add_heading(doc, "7. SUMMARY")
    add_body(
        doc,
        "The Gym Management System successfully demonstrates the design and implementation of a database-backed management application. It integrates a MySQL schema with a Node.js REST API and a browser-based frontend to support member management, trainer management, package enrollment, payments, and reporting.",
    )
    add_body(
        doc,
        "The project applies core DBMS concepts in a practical setting and provides a maintainable foundation for future improvements such as attendance tracking, automated renewal notifications, member self-service features, and advanced analytics.",
    )

    add_heading(doc, "8. REFERENCES")
    refs = [
        "Ramez Elmasri and Shamkant B. Navathe, Fundamentals of Database Systems, Pearson.",
        "MySQL 8.0 Reference Manual, Oracle Corporation.",
        "Express.js Documentation, OpenJS Foundation.",
        "Node.js Documentation, OpenJS Foundation.",
        "Bootstrap and Chart.js documentation for frontend layout and visualization references.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(ref)
        style_run(r, size=10.5)

    add_heading(doc, "APPENDIX A: PROJECT STRUCTURE")
    add_code_block(
        doc,
        """gym management system/
  database/schema.sql
  backend/server.js
  backend/config/db.js
  backend/controllers/
  backend/routes/
  backend/middleware/auth.js
  frontend/index.html
  frontend/style.css
  frontend/app.js""",
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Gym Management System | Database Management Systems Laboratory"
        for run in footer.runs:
            style_run(run, size=8, color="6B7280")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
